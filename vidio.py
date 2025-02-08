from docx import Document
# 获取当前脚本所在目录的绝对路径
import os

# 获取Python解释器（或exe）所在目录
exe_dir = os.path.dirname(sys.executable)
print(exe_dir)
script_dir = os.path.dirname(sys.executable)
print(script_dir)
# 改变当前工作目录到exe文件所在的目录
os.chdir(exe_dir)
source_dir = os.path.dirname(sys.executable)
print(f"正确工作路径 directory: {os.getcwd()}")
# 需要检查和创建的文件列表
doc_files = ['input.docx','moxing.docx','input1.docx', 'output.docx', 'temp.docx', 'text3.docx','定位编辑.docx','fixtext.docx', 'text2.docx']
# 检查每个文件是否存在，如果不存在则创建一个空文档
for filename in doc_files:
    file_path = os.path.join(script_dir, filename)

    # 如果文件不存在，则创建一个新的空Word文档
    if not os.path.exists(file_path):
        try:
            doc = Document()
            doc.save(file_path)
            print(f"已创建文件: {filename}")
        except Exception as e:
            print(f"无法创建文件{filename}：{e}")
    else:
        print(f"文件 {filename} 已存在，无需操作。")
def read_replacement_rules_from_doc(file_path, delimiter=':'):
    # 打开Word文档
    doc = Document(file_path)

    # 初始化一个空字典存放替换规则
    replacement_rules = {}

    # 遍历所有段落
    for paragraph in doc.paragraphs:
        text = paragraph.text
        parts = text.split(delimiter)
        if len(parts) == 2 and parts[0].strip() and parts[1].strip():
            key = parts[0].strip()
            value = parts[1].strip()
            replacement_rules[key] = value

    return replacement_rules

# 从fixtext.docx中读取替换规则
replacement_rules = read_replacement_rules_from_doc('fixtext.docx')

# 打印读取到的替换规则
print(replacement_rules)

# 检查并打开docx文件
doc_path = os.path.join(script_dir, 'input1.docx')
if not os.path.exists(doc_path):
    raise FileNotFoundError("找不到文件：input1.docx")

doc = Document(doc_path)

# 遍历所有段落进行替换
for paragraph in doc.paragraphs:
    for old_word, new_word in replacement_rules.items():
        paragraph.text = paragraph.text.replace(old_word, new_word)

# 将修改后的内容保存回原文件
doc.save(doc_path)

def remove_empty_paragraphs(file_path):
    # 打开Word文档
    doc = Document(file_path)

    # 遍历所有段落
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        if len(paragraph.text.strip()) == 0:  # 检查段落文本是否为空（考虑可能有空白字符）
            paragraphs_to_remove.append(paragraph)

    # 删除空段落（不直接在遍历中删除以避免迭代器错误）
    for paragraph in paragraphs_to_remove:
        paragraph._element.getparent().remove(paragraph._element)

    # 保存修改后的文档
    doc.save(file_path)

# 使用函数处理文件
remove_empty_paragraphs('input1.docx')
# 获取当前脚本所在目录



# 定义替换规则
replacement_rules = {'论文': '洛文',}

# 检查并打开docx文件
doc_path = os.path.join(script_dir, 'input1.docx')
if not os.path.exists(doc_path):
    raise FileNotFoundError("找不到文件：input1.docx")

doc = Document(doc_path)

# 遍历所有段落进行替换
for paragraph in doc.paragraphs:
    for old_word, new_word in replacement_rules.items():
        paragraph.text = paragraph.text.replace(old_word, new_word)

# 将修改后的内容保存回原文件
doc.save(doc_path)
# 现在所有指定的文件都应在当前目录下存在，可以进行后续读取或写入操作。

# 现在所有指定的文件都应在当前目录下存在，可以进行后续读取或写入操作。
import shutil

# 复制文件
shutil.copyfile('output.docx', 'temp.docx')

# 然后使用以下代码清空并保存 'temp.docx' 为 'output.docx'
from docx import Document

# 打开文档
doc1 = Document('temp.docx')

# 清空文档内容并删除空段落
for paragraph in doc1.paragraphs:
    if not paragraph.text.strip():  # 如果段落文本为空或仅包含空格
        paragraph._element.getparent().remove(paragraph._element)  # 删除该段落元素
    else:
        paragraph.clear()

# 保存更改为 'output.docx'
doc1.save('output.docx')

os.environ['DASHSCOPE_API_KEY'] = 'sk-ade26912d9f6406fabe8edd7c5b2b7b1'
from http import HTTPStatus
import dashscope
import json
from http import HTTPStatus
import dashscope
from docx import Document


def merge_two_docs(doc1_path, doc2):
    # 读取第一个文档
    doc1 = Document(doc1_path)

    # 将第二个文档的内容添加到第一个文档的末尾
    for paragraph in doc2.paragraphs:
        doc1.add_paragraph(paragraph.text)

    # 保存合并后的文档
    doc1.save(doc1_path)


def save_to_docx(response):
    doc = Document()

    for choice in response.output['choices']:
        content = choice['message']['content']
        content += 'A'  # 在内容末尾添加大写的'A'
        doc.add_paragraph(content)

    # 合并文档
    merge_two_docs('output.docx', doc)
    print("内容已添加到output.docx")


def read_text_from_doc(file_path, batch_size=500, min_batch_size=500, setup_info="text2 "):
    doc = Document(file_path)
    text_batches = []
    current_batch = ""
    text2 = ""
    for paragraph in doc.paragraphs:
        current_paragraph_text = paragraph.text.strip()
        if len(current_batch + current_paragraph_text) <= batch_size:
            current_batch += f"{setup_info}{current_paragraph_text}\n"

        else:
            # 如果当前段落使总字符数超过batch_size，则将当前批次添加到text_batches并开始新的批次
            text_batches.append(current_batch)
            current_batch = f"{setup_info}{current_paragraph_text}\n"


    # 处理最后一批次，如果小于min_batch_size，则将其与前一个批次合并
    if len(current_batch) >= min_batch_size:

        text_batches.append(current_batch)
    elif len(text_batches) > 0:
        text_batches[-1] += current_batch
    # 获取用户输入
    



    user_input = None

    user_input = 1
    print("你选择了数字1")


    # 读取文档text3和text2
    def read_docx(filename):
        document = Document(filename)
        text_content = [paragraph.text for paragraph in document.paragraphs]
        return '\n'.join(text_content)

    # 获取当前脚本所在目录

    # 计算文本文件的相对路径
    file_path_text3 = os.path.join(script_dir, 'text3.docx')
    file_path_text2 = os.path.join(script_dir, 'text2.docx')

    text3 = read_docx(file_path_text3)
    text2 = read_docx(file_path_text2)

    print("text3的内容：")
    print(text3)

    print("\ntext2的内容：")
    print(text2)

    text0 = ""
    # 根据用户输入给text0赋值
    if user_input == 1:
        text2 = ""
    elif user_input == 2:
        text3 = ""
    else:
        # 如果输入既不是1也不是2，则可以给出错误提示或其他处理方式
        print("无效输入，请输入1或2")

    for i in range(len(text_batches)):
        text_batches[i] = text3 + "" + text_batches[i] + "" + text2

    # for i in range(len(text_batches)):
    #  text_batches[i] = text_batches[i] + "" + text2
    return text_batches


input_file_path = 'input1.docx'



# 确保ai处理的函数独立
from openai import OpenAI

from docx import Document

def culi(a, api_key,fieldQ):
    # 创建 OpenAI 客户端实例
    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
    
    # 使用传入的消息列表a进行聊天
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=a
    )
    
    # 将API的响应添加到消息列表中
    
    a = response.choices[0].message
    result_string = str(a)
    content_start = result_string.find('content=')
    if content_start != -1:
        content_start += len('content=')
        content_end = result_string.find(', refusal=None', content_start)
        if content_end == -1:
            content_end = len(result_string)
        content = result_string[content_start:content_end].strip('"')


    sd_content = content.strip("'")
    sd_content = sd_content.replace('\\n', '\n')
    sd_content = sd_content + "\nA"
    print(sd_content)
    doc = Document('output.docx')
    chinese_punctuation = "，。……！？…………；：、（）〈〉《》{}【】“”‘’"
    from fuzzywuzzy import process
    # 获取所有可能的匹配项，按分数排序
    import re
    
    sentences = re.split(r'([' + re.escape(chinese_punctuation) + '])', sd_content)
    # 将标点符号重新拼接到句子上
    sentences = [sentences[i] + sentences[i+1] for i in range(0, len(sentences)-1, 2)]
    pattern = re.compile(r".*[" + re.escape(chinese_punctuation) + "]$")
    candidates = [s for s in sentences if pattern.match(s)]
    matches = process.extract(fieldQ, candidates, limit=20)
    
# 遍历匹配项，找到 matched_word
    matched_word = None
    score = 0
    for match, match_score in matches:
        if match and match[-1] in chinese_punctuation and match_score > pipeisuzi:
           matched_word = match
           score = match_score
           break
           
    if matched_word:
       position = sd_content.find(matched_word)
       if position != -1:
           sd_content = sd_content[:position + len(matched_word)]
    else:
        sd_content = sd_content
    doc.add_paragraph(sd_content)

    # 保存修改后的文档
    doc.save('output.docx')






# 读取Word文档
doc = Document('AIapi.docx')

gpttext = ''

# 提取段落文本
for paragraph in doc.paragraphs:
    gpttext += paragraph.text  # 保留段落换行

# 使用示例
api_key = gpttext  # 应该从安全的地方获取



# 打开 .docx 文件
doc = Document('suzi.docx')

# 初始化一个空字符串来存储文档内容
content = ''

# 遍历文档中的每个段落，并将其内容添加到变量 content 中
for para in doc.paragraphs:
    content += para.text.strip()  # 使用 strip() 去除空白字符

# 假设 content 是一个数字，将其转换为整数或浮点数
try:
    if '.' in content:  # 如果包含小数点，转换为浮点数
        suzi = float(content)
        
    else:  # 否则转换为整数
        suzi = int(content)
        
except ValueError:
    print("文件内容不是一个有效的数字！")
    suzi = None  # 如果转换失败，将 a 设置为 None


# 打开 .docx 文件
doc = Document('pipeisuzi.docx')

# 初始化一个空字符串来存储文档内容
content = ''

# 遍历文档中的每个段落，并将其内容添加到变量 content 中
for para in doc.paragraphs:
    content += para.text.strip()  # 使用 strip() 去除空白字符

# 假设 content 是一个数字，将其转换为整数或浮点数
try:
    if '.' in content:  # 如果包含小数点，转换为浮点数
        pipeisuzi = float(content)
        
    else:  # 否则转换为整数
        pipeisuzi = int(content)
        
except ValueError:
    print("文件内容不是一个有效的数字！")
    pipeisuzi = None  # 如果转换失败，将 a 设置为 None



if __name__ == '__main__':
    text2 = ""
    text_batches = read_text_from_doc(input_file_path, setup_info=text2)
    for i, text_batch in enumerate(text_batches):
        fieldQ = text_batch[-suzi:]
        messages = [{"role": "user", "content": text_batch}]
        #culi(messages, api_key,fieldQ)
from docx import Document
    
# 使用您的实际文件路径替换 'your_file_path.docx'

# 获取当前脚本所在目录的绝对路径


# 构建output.docx文件的绝对路径
file_path = os.path.join(script_dir, 'output.docx')
print(file_path)
print("位置在于:")


print(source_dir)
def remove_empty_paragraphs(doc_path):
    doc = Document(doc_path)
    paragraphs = doc.paragraphs
    for para in paragraphs:
        if not para.text.strip():
            p = para._element
            p.getparent().remove(p)
    doc.save(doc_path)

# 调用函数，替换成你的文档路径
remove_empty_paragraphs('output.docx')


#读取对应参数

# 其余原有导入保持不变 
# 其余原有导入保持不变 
import numpy as np 

from moviepy.editor import *
from moviepy.config import change_settings 

def load_parameters(doc_path):
    doc = Document(doc_path)
    params = {"background": {}, "dialog": {}, "text": {}, "output": {}}
    current_section = None  # 类型转换规则 
    converters = {
        "int": int, 
        "float": float,
        "rgb": lambda x: tuple(map(int, x.split(','))),
        "bool": lambda x: x.lower() == "true"
    }
    
    # 参数类型映射表 
    TYPE_MAP = {
        "background": {
            "default_duration": "float",
            "resolution": "optional"
        },
        "dialog": {
            "width_ratio": "float", 
            "height_ratio": "float",
            "position_y": "float", 
            "bg_alpha": "float",
            "border_size": "int", 
            "border_radius": "int"
        },
        "text": {
            "size": "int", 
            "speed": "int",
            "padding_x": "int", 
            "padding_y": "int",
            "line_spacing": "float"
        },
        "output": {
            "fps": "int", 
            "threads": "int",
            "audio_enabled": "bool"
        }
    }

    for para in doc.paragraphs:
        line = para.text.strip()
        if not line or line.startswith("#"):
            continue 
        
        # 识别段落分类 
        if line.startswith("[") and line.endswith("]"):
            current_section = line[1:-1].lower()
            continue
            
        if "=" in line and current_section:
            key, value = map(str.strip, line.split("=", 1))
            param_type = TYPE_MAP.get(current_section, {}).get(key, "str")
            
            try:
                if param_type in converters:
                    params[current_section][key] = converters[param_type](value)
                elif "_color" in key:
                    params[current_section][key] = converters["rgb"](value)
                else:
                    params[current_section][key] = value 
            except:
                print(f"参数解析失败：{current_section}.{key} = {value}")
                params[current_section][key] = value 
                
    return params 

def generate_video(script_dir):  # 🆕 修改1：添加参数 
    # 🆕 修改2：参数文件路径锚定 
    params = load_parameters(os.path.join(script_dir, "Parameter.docx"))
    
    # 🆕 修改3：背景路径锚定 
    bg_path = os.path.join(script_dir, params["background"]["background_path"])
    
    if bg_path.lower().endswith(('.png', '.jpg', '.jpeg')):
        bg_clip = ImageClip(bg_path).set_duration(
            params["background"].get("default_duration", 10)
        )
    else:
        bg_clip = VideoFileClip(bg_path)
    
    # 分辨率处理 
    if "resolution" in params["background"]:
        w, h = map(int, params["background"]["resolution"].split('x'))
        bg_clip = bg_clip.resize((w, h))
    
    # 创建动态对话框 
    dialog_w = int(bg_clip.w * params["dialog"]["width_ratio"])
    dialog_h = int(bg_clip.h * params["dialog"]["height_ratio"])
    
    def create_dialog(t):
        return (ColorClip(size=(dialog_w, dialog_h), color=params["dialog"]["bg_color"])
                .set_opacity(params["dialog"]["bg_alpha"])
                .set_position(('center', bg_clip.h * params["dialog"]["position_y"]))
                .margin(
                    top=params["text"]["padding_y"], 
                    bottom=params["text"]["padding_y"],
                    left=params["text"]["padding_x"],
                    right=params["text"]["padding_x"],
                    color=params["dialog"]["border_color"]
                )
                .set_duration(bg_clip.duration))
    
    # 文字动画生成 
    text_content = params["text"]["content"]
    def text_animation(t):
        chars_show = min(int(t * params["text"]["speed"]), len(text_content))
        current_text = text_content[:chars_show]
        
        return (TextClip(
            txt=current_text,
            font=params["text"]["font"],
            fontsize=params["text"]["size"],
            color=params["text"]["color"],
            align='west',
            size=(dialog_w - 2*params["text"]["padding_x"], None),
            method='caption',
            print_cmd=True 
        )
        .set_position((
            params["text"]["padding_x"], 
            params["text"]["padding_y"]
        )))
    
    # 合成最终视频 
    final_clip = CompositeVideoClip([
        bg_clip,
        create_dialog(0).crossfadein(0.5),
        text_animation(0).set_start(0.5)
    ], use_bgclip=True).set_duration(bg_clip.duration)
    
    # 音频处理 
    if params["output"].get("audio_enabled", True) and hasattr(bg_clip, 'audio'):
        final_clip = final_clip.set_audio(bg_clip.audio)
    
    # 🆕 修改4：输出路径锚定 
    output_path = os.path.join(script_dir, params["output"]["path"])
    final_clip.write_videofile(
        output_path, # 使用锚定后的路径 
        fps=params["output"]["fps"],
        codec=params["output"]["codec"],
        threads=params["output"]["threads"],
        preset='slow',
        audio_codec='aac' if params["output"]["audio_enabled"] else None 
    )

if __name__ == "__main__":
    # 🆕 修改5：传递已定义的根目录变量 
    script_dir = os.path.dirname(os.path.abspath(__file__))  # 示例定义（实际由用户定义）
    generate_video(script_dir)